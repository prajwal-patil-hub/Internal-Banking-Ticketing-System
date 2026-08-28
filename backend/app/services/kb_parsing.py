"""Turn an uploaded file into text, preserving the structure retrieval needs.

Parsing is the stage that quietly decides how good retrieval can ever be. Two
things have to survive the trip out of the file, and both are easy to lose:

**Headings.** A passage that reads "must be raised within 45 days" is useless
without "3.2 Chargeback timelines" above it. Headings are emitted as separate
`Block`s so the chunker can carry the active heading path into every chunk it
produces — the retrieval index then contains the context, not just the clause.

**Tables.** A chargeback matrix flattened into running prose is worse than
useless: it reads as fluent English and is factually scrambled, so the model
cites it confidently. Tables are emitted as pipe-delimited rows, one row per
line, and the chunker is told never to split inside one.

Scanned PDFs are out of scope: this extracts the *text layer*, and a scan has
none. `extract()` reports how many pages produced no text so the caller can
tell the operator "this looks like a scan, OCR is not enabled" rather than
silently indexing an empty document.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

from app.core.exceptions import ValidationError
from app.core.logging import get_logger

log = get_logger(__name__)

#: Extensions this module can turn into text. Anything else is refused at
#: upload time rather than stored and found unparseable later.
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({"pdf", "docx", "txt", "md", "csv"})

#: Largest total text a single document may expand to.
#:
#: DOCX is a zip, and deflate reaches roughly 1000:1 on repetitive XML, so a
#: 40 MB upload that passes the size check can expand to tens of gigabytes
#: while python-docx reads each part fully into memory. The upload limit
#: bounds what arrives; this bounds what it becomes.
MAX_EXPANDED_CHARS = 20 * 1024 * 1024

#: Pages read from one PDF. Text extraction is synchronous and in-request, so
#: an unbounded page count is a slow-loris that needs no malformed input.
MAX_PDF_PAGES = 2000


@dataclass
class Block:
    """One structural unit of the document.

    `kind` drives chunking policy, not presentation:
      - "heading" starts a new section and updates the heading path
      - "table_row" must never be split away from its neighbours
      - "text" is ordinary prose and may be split on sentence boundaries
    """

    kind: str
    text: str
    level: int = 0
    page: int | None = None


@dataclass
class ParsedDocument:
    blocks: list[Block]
    page_count: int | None = None
    #: Pages whose text layer was empty. A high count against a PDF means a
    #: scan, which this pipeline cannot read.
    empty_pages: int = 0


def parse(data: bytes, extension: str, filename: str = "") -> ParsedDocument:
    """Dispatch on extension. Raises ValidationError for unsupported types."""
    ext = (extension or "").lower().lstrip(".")
    if ext == "pdf":
        return _parse_pdf(data)
    if ext == "docx":
        return _parse_docx(data)
    if ext in {"txt", "md", "csv"}:
        return _parse_text(data, ext)
    raise ValidationError(
        f"'{ext or 'unknown'}' files cannot be indexed. Supported: "
        f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}."
    )


# ---------------------------------------------------------------------------
# Plain text and Markdown
# ---------------------------------------------------------------------------

#: ATX headings (`## Title`) and the common numbered-clause style used in bank
#: policy documents (`3.2 Chargeback timelines`). The second pattern is
#: deliberately strict — it needs the number, a space and a capital — because a
#: looser rule promotes ordinary sentences beginning with a figure into
#: headings, which then poison the heading path of everything beneath them.
_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")
_NUMBERED_HEADING = re.compile(r"^((?:\d+\.){0,4}\d+)[.)]?\s+([A-Z][^.!?]{2,80})$")


def _decode(data: bytes) -> str:
    """UTF-8, falling back to latin-1 rather than throwing.

    A policy PDF exported by an old Windows tool is routinely cp1252. Refusing
    the upload over one smart quote helps nobody; latin-1 decodes every byte,
    so the worst case is a mangled character rather than a lost document.
    """
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")


def _parse_text(data: bytes, ext: str) -> ParsedDocument:
    text = _decode(data)
    blocks: list[Block] = []

    if ext == "csv":
        # Every CSV line is a table row: never split one, and keep the header
        # adjacent so a retrieved fragment still has column names.
        for line in text.splitlines():
            if line.strip():
                blocks.append(Block(kind="table_row", text=line.strip()))
        return ParsedDocument(blocks=blocks)

    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue

        md = _MD_HEADING.match(line)
        if md:
            blocks.append(Block(kind="heading", text=md.group(2), level=len(md.group(1))))
            continue

        num = _NUMBERED_HEADING.match(line.strip())
        if num:
            level = num.group(1).count(".") + 1
            blocks.append(
                Block(kind="heading", text=f"{num.group(1)} {num.group(2)}", level=level)
            )
            continue

        # A markdown table row. Kept whole for the reason in the module docstring.
        if line.lstrip().startswith("|") and line.rstrip().endswith("|"):
            blocks.append(Block(kind="table_row", text=line.strip()))
            continue

        blocks.append(Block(kind="text", text=line.strip()))

    return ParsedDocument(blocks=blocks)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _parse_pdf(data: bytes) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise ValidationError("PDF support requires the `pypdf` package.") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ValidationError(f"This PDF could not be opened: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        # Some PDFs are "encrypted" with an empty owner password and open fine.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ValidationError(
                "This PDF is password-protected and cannot be indexed. "
                "Upload an unprotected copy."
            ) from exc

    page_count = len(reader.pages)
    if page_count > MAX_PDF_PAGES:
        raise ValidationError(
            f"This PDF has {page_count:,} pages, over the {MAX_PDF_PAGES:,} "
            "limit. Text extraction runs inside the request, so a document "
            "this long would stall the API. Split it into parts."
        )

    blocks: list[Block] = []
    empty_pages = 0
    extracted_chars = 0

    for page_no, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            log.warning("kb.pdf_page_failed", page=page_no, error=str(exc))
            text = ""

        if not text.strip():
            empty_pages += 1
            continue

        extracted_chars += len(text)
        if extracted_chars > MAX_EXPANDED_CHARS:
            raise ValidationError(
                f"This PDF expands to more than "
                f"{MAX_EXPANDED_CHARS // 1_048_576} MB of text, which is more "
                "than the knowledge base will index from one document."
            )

        for raw in text.splitlines():
            line = raw.strip()
            if not line:
                continue
            num = _NUMBERED_HEADING.match(line)
            if num:
                level = num.group(1).count(".") + 1
                blocks.append(
                    Block(
                        kind="heading",
                        text=f"{num.group(1)} {num.group(2)}",
                        level=level,
                        page=page_no,
                    )
                )
            else:
                blocks.append(Block(kind="text", text=line, page=page_no))

    # Every page empty means a scan (or a broken text layer). Indexing that
    # produces a document that exists, retrieves nothing, and looks like a bug
    # in retrieval rather than a bad upload.
    if page_count and empty_pages == page_count:
        raise ValidationError(
            "No text could be extracted from this PDF — it appears to be a "
            "scanned image. OCR is not enabled, so upload a text-based PDF."
        )

    return ParsedDocument(blocks=blocks, page_count=page_count, empty_pages=empty_pages)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _reject_zip_bomb(data: bytes) -> None:
    """Refuse an archive whose declared expansion is absurd.

    Reads only the central directory: `file_size` is metadata, so nothing is
    decompressed to make this decision. Both an absolute cap and a ratio cap
    apply — a small file with a huge ratio and a large file with a modest one
    are both ways to arrive at the same out-of-memory.
    """
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            total = sum(info.file_size for info in archive.infolist())
    except zipfile.BadZipFile as exc:
        raise ValidationError(
            "This file is not a readable Word document."
        ) from exc

    if total > MAX_EXPANDED_CHARS:
        raise ValidationError(
            f"This document expands to {total / 1_048_576:.0f} MB, over the "
            f"{MAX_EXPANDED_CHARS // 1_048_576} MB limit. Split it into parts."
        )
    if data and total / len(data) > 200:
        raise ValidationError(
            "This document's compression ratio is implausible for a Word file "
            "and it has been refused as a safety measure."
        )


def _parse_docx(data: bytes) -> ParsedDocument:
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise ValidationError("DOCX support requires the `python-docx` package.") from exc

    # Inspect the zip directory before handing the bytes to python-docx, which
    # reads each part fully into memory. The directory declares uncompressed
    # sizes, so the bomb is detectable without decompressing anything.
    _reject_zip_bomb(data)

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ValidationError(f"This Word document could not be opened: {exc}") from exc

    blocks: list[Block] = []

    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue

        # Word's own heading styles are the most reliable structure signal in
        # a docx — far better than guessing from the text, so prefer them.
        style_name = (getattr(para.style, "name", "") or "").lower()
        if style_name.startswith("heading"):
            digits = "".join(ch for ch in style_name if ch.isdigit())
            blocks.append(Block(kind="heading", text=text, level=int(digits or 1)))
            continue

        num = _NUMBERED_HEADING.match(text)
        if num:
            blocks.append(
                Block(
                    kind="heading",
                    text=f"{num.group(1)} {num.group(2)}",
                    level=num.group(1).count(".") + 1,
                )
            )
            continue

        blocks.append(Block(kind="text", text=text))

    # python-docx keeps tables out of `paragraphs`, so they must be walked
    # separately or every table in the document silently vanishes.
    for table in document.tables:
        for row in table.rows:
            cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                blocks.append(Block(kind="table_row", text=" | ".join(cells)))

    return ParsedDocument(blocks=blocks)
