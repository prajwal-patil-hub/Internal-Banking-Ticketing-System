"""Parsing and chunking — the stage that decides how good retrieval can get.

These tests generate real PDFs and DOCX files rather than asserting against
fixtures, because the failures worth catching here are format failures: a
table that python-docx keeps outside `paragraphs`, a PDF whose text layer is
empty. A hand-written fixture would encode my assumption about the format
rather than the format itself.
"""

from __future__ import annotations

import io

import pytest

from app.core.exceptions import ValidationError
from app.services.kb_chunking import chunk_blocks
from app.services.kb_parsing import parse


def _pdf(pages: list[list[str]]) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    for lines in pages:
        y = 780
        for line in lines:
            c.drawString(72, y, line)
            y -= 20
        c.showPage()
    c.save()
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Structure survives parsing
# ---------------------------------------------------------------------------

def test_markdown_headings_nest_into_a_path() -> None:
    doc = parse(
        b"# Policy\n\n## 3. Chargebacks\n\n### 3.2 Timelines\n\nRaise within 45 days.\n",
        "md",
    )
    chunks = chunk_blocks(doc.blocks, max_chars=500, overlap_chars=50)
    assert len(chunks) == 1
    assert chunks[0].heading_path == "Policy > 3. Chargebacks > 3.2 Timelines"


def test_heading_path_is_embedded_not_just_stored() -> None:
    """A passage retrieved alone must carry its section into the vector.

    Storing the path but embedding only the body would mean "within 45 days"
    is indexed with no indication of what it applies to.
    """
    doc = parse(b"## 3.2 Timelines\n\nRaise within 45 days.\n", "md")
    chunk = chunk_blocks(doc.blocks, max_chars=500, overlap_chars=50)[0]
    assert "3.2 Timelines" in chunk.embedding_text()
    assert chunk.embedding_text() != chunk.content


def test_a_new_heading_starts_a_new_chunk() -> None:
    doc = parse(
        b"## A\n\nAlpha text here.\n\n## B\n\nBravo text here.\n",
        "md",
    )
    chunks = chunk_blocks(doc.blocks, max_chars=5000, overlap_chars=0)
    # Both fit inside one chunk by size; the heading is what separates them.
    assert len(chunks) == 2
    assert chunks[0].heading_path == "A"
    assert chunks[1].heading_path == "B"


def test_table_rows_are_never_merged_into_prose() -> None:
    doc = parse(
        b"## Windows\n\nSome prose about windows.\n\n| Type | Days |\n| Fraud | 120 |\n",
        "md",
    )
    chunks = chunk_blocks(doc.blocks, max_chars=5000, overlap_chars=0)
    kinds = [sorted(c.kinds) for c in chunks]
    assert ["text"] in kinds
    assert ["table_row"] in kinds
    # No chunk mixes the two.
    assert not any(len(k) > 1 for k in kinds)


def test_ordinals_are_dense_and_unique() -> None:
    """(version_id, ordinal) is a unique constraint; a gap or repeat breaks insert."""
    body = "\n\n".join(f"## S{i}\n\nBody {i} " + ("x" * 200) for i in range(6))
    doc = parse(body.encode(), "md")
    chunks = chunk_blocks(doc.blocks, max_chars=300, overlap_chars=40)
    ordinals = [c.ordinal for c in chunks]
    assert ordinals == list(range(len(chunks)))


def test_oversized_prose_splits_on_sentences_with_overlap() -> None:
    text = " ".join(f"Sentence number {i} carries a fact." for i in range(60))
    doc = parse(f"## S\n\n{text}\n".encode(), "md")
    chunks = chunk_blocks(doc.blocks, max_chars=300, overlap_chars=60)
    assert len(chunks) > 1
    # Overlap means consecutive chunks share some text.
    assert any(
        chunks[i].content[-30:] in chunks[i + 1].content
        or chunks[i + 1].content[:30] in chunks[i].content
        for i in range(len(chunks) - 1)
    )


def test_a_single_sentence_longer_than_the_limit_is_hard_cut() -> None:
    """Otherwise one enormous chunk dominates every retrieval it appears in."""
    doc = parse(("## S\n\n" + "x" * 900 + "\n").encode(), "md")
    chunks = chunk_blocks(doc.blocks, max_chars=200, overlap_chars=20)
    assert len(chunks) > 1
    assert all(c.char_count <= 200 for c in chunks)


# ---------------------------------------------------------------------------
# Formats
# ---------------------------------------------------------------------------

def test_pdf_keeps_page_numbers_for_citation_deep_links() -> None:
    doc = parse(_pdf([["3.2 Timelines", "Raise within 45 days."], ["3.3 Evidence", "Attach the form."]]), "pdf")
    assert doc.page_count == 2
    pages = {b.page for b in doc.blocks}
    assert pages == {1, 2}


def test_scanned_pdf_is_refused_not_indexed_empty() -> None:
    """A scan has no text layer. Indexing it creates a document that exists,
    retrieves nothing, and looks like a retrieval bug rather than a bad upload."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.rect(100, 100, 200, 200, fill=1)
    c.showPage()
    c.save()

    with pytest.raises(ValidationError, match="scanned"):
        parse(buf.getvalue(), "pdf")


def test_docx_tables_are_extracted() -> None:
    """python-docx keeps tables out of `paragraphs` — walking only paragraphs
    silently drops every table in the document."""
    import docx

    d = docx.Document()
    d.add_heading("Chargebacks", level=1)
    d.add_paragraph("Prose about chargebacks.")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Fraud"
    t.rows[0].cells[1].text = "120 days"
    buf = io.BytesIO()
    d.save(buf)

    doc = parse(buf.getvalue(), "docx")
    table_rows = [b for b in doc.blocks if b.kind == "table_row"]
    assert table_rows and "120 days" in table_rows[0].text


def test_docx_heading_styles_beat_text_guessing() -> None:
    import docx

    d = docx.Document()
    d.add_heading("Compliance", level=1)
    d.add_paragraph("Body text.")
    buf = io.BytesIO()
    d.save(buf)

    doc = parse(buf.getvalue(), "docx")
    assert doc.blocks[0].kind == "heading"
    assert doc.blocks[0].level == 1


def test_csv_rows_stay_whole() -> None:
    doc = parse(b"type,window\nfraud,120\nservice,45\n", "csv")
    assert all(b.kind == "table_row" for b in doc.blocks)


def test_unsupported_extension_is_refused() -> None:
    with pytest.raises(ValidationError):
        parse(b"\x00\x01", "xlsx")


def test_non_utf8_bytes_do_not_kill_the_upload() -> None:
    """A cp1252 export shouldn't be rejected over one smart quote."""
    # U+201C/U+201D are the smart quotes; cp1252 encodes them as 0x93/0x94,
    # which are invalid UTF-8 and would raise on a strict decode.
    raw = "## S\n\nRs. 500 “quoted” amount.\n".encode("cp1252")
    assert b"\x93" in raw
    doc = parse(raw, "md")
    assert any("500" in b.text for b in doc.blocks)


def test_numbered_heading_detection_does_not_promote_ordinary_sentences() -> None:
    """A loose rule turns any sentence starting with a figure into a heading,
    which then poisons the heading path of everything beneath it."""
    doc = parse(b"## Real\n\n45 days is the limit for service disputes and related claims.\n", "md")
    headings = [b.text for b in doc.blocks if b.kind == "heading"]
    assert headings == ["Real"]
