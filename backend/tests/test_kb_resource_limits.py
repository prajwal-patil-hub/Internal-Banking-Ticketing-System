"""Resource exhaustion: oversized bodies, zip bombs, and runaway page counts.

These are the two review findings that survived the first fix pass, because
both needed real machinery rather than a tighter conditional. The body limit
had to move above dependency resolution — FastAPI parses a multipart body
before it solves dependencies, so the auth check is not early enough to be the
gate. The expansion limits had to move before python-docx reads the archive.

Everything here builds a genuinely hostile input rather than asserting on a
constant, because the interesting question is not "is the number right" but
"does the check run before the damage".
"""

from __future__ import annotations

import io
import zipfile

import pytest
from fastapi import FastAPI, Request
from fastapi.testclient import TestClient

from app.core.exceptions import ValidationError
from app.middleware.body_limit import BodySizeLimitMiddleware
from app.services.kb_parsing import MAX_PDF_PAGES, parse

# ---------------------------------------------------------------------------
# Body size, enforced above the application
# ---------------------------------------------------------------------------

def _app(default_limit: int = 1024, limits: dict[str, int] | None = None) -> TestClient:
    app = FastAPI()

    # Raw-body endpoints on purpose: a typed body would make FastAPI reject
    # "xxxx" as malformed JSON with a 400, and the assertions would then be
    # measuring the request parser rather than the middleware under test.
    @app.post("/api/v1/echo")
    async def echo(request: Request) -> dict:  # pragma: no cover - trivial
        return {"bytes": len(await request.body())}

    @app.post("/api/v1/kb/upload")
    async def upload(request: Request) -> dict:  # pragma: no cover - trivial
        return {"bytes": len(await request.body())}

    app.add_middleware(
        BodySizeLimitMiddleware, default_limit=default_limit, limits=limits or {}
    )
    return TestClient(app)


def test_oversized_body_is_refused_with_413() -> None:
    client = _app(default_limit=1024)
    resp = client.post("/api/v1/echo", content=b"x" * 4096)
    assert resp.status_code == 413
    assert resp.json()["error"]["code"] == "payload_too_large"


def test_body_within_the_limit_passes_through() -> None:
    """And the app receives the body intact — the middleware drains the stream
    to measure it, so it has to replay every byte it consumed."""
    client = _app(default_limit=8192)
    resp = client.post("/api/v1/echo", content=b"x" * 4096)
    assert resp.status_code == 200
    assert resp.json()["bytes"] == 4096


def test_the_knowledge_base_gets_its_own_wider_limit() -> None:
    """A single global cap would either starve the KB or widen every other
    endpoint to no purpose."""
    client = _app(default_limit=1024, limits={"/api/v1/kb/": 65536})
    assert client.post("/api/v1/kb/upload", content=b"x" * 4096).status_code == 200
    assert client.post("/api/v1/echo", content=b"x" * 4096).status_code == 413


def test_a_lying_content_length_does_not_get_through() -> None:
    """The cheap header check is bypassable; the streaming count is not.

    Sending a generator body makes httpx use chunked transfer encoding, so no
    Content-Length is declared at all — which is exactly the bypass.
    """
    client = _app(default_limit=1024)

    def chunks():
        for _ in range(10):
            yield b"x" * 512

    resp = client.post("/api/v1/echo", content=chunks())
    assert resp.status_code == 413


def test_reads_are_not_penalised() -> None:
    client = _app(default_limit=1)
    # GET carries no body; the middleware must not touch it.
    assert client.get("/docs").status_code in (200, 404)


# ---------------------------------------------------------------------------
# Decompression bombs
# ---------------------------------------------------------------------------

def _zip_bomb(uncompressed_mb: int) -> bytes:
    """A real archive whose directory declares a huge expansion.

    Highly repetitive content, so deflate produces a tiny file — the same
    shape as a malicious DOCX.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("word/document.xml", b"\0" * (uncompressed_mb * 1024 * 1024))
    return buf.getvalue()


def test_a_zip_bomb_is_refused_before_it_is_decompressed() -> None:
    bomb = _zip_bomb(64)
    # The upload itself is small — it passes every size check on the way in.
    assert len(bomb) < 1024 * 1024

    with pytest.raises(ValidationError) as exc:
        parse(bomb, "docx")
    assert "expands" in str(exc.value) or "compression ratio" in str(exc.value)


def test_a_small_archive_with_an_absurd_ratio_is_refused() -> None:
    """Below the absolute cap but still a 1000:1 expansion — the ratio check
    is what catches this one."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("word/document.xml", b"\0" * (8 * 1024 * 1024))
    data = buf.getvalue()

    with pytest.raises(ValidationError):
        parse(data, "docx")


def test_a_normal_docx_still_parses() -> None:
    """The guard must not refuse legitimate documents."""
    import docx

    d = docx.Document()
    d.add_heading("Chargebacks", level=1)
    d.add_paragraph("Raise within 45 days of the transaction date.")
    buf = io.BytesIO()
    d.save(buf)

    parsed = parse(buf.getvalue(), "docx")
    assert any("45 days" in b.text for b in parsed.blocks)


def test_a_file_that_is_not_a_zip_is_refused_clearly() -> None:
    with pytest.raises(ValidationError, match="not a readable Word document"):
        parse(b"this is plainly not a docx", "docx")


# ---------------------------------------------------------------------------
# PDF page count
# ---------------------------------------------------------------------------

def test_a_pdf_over_the_page_cap_is_refused() -> None:
    """Text extraction is synchronous and in-request, so an unbounded page
    count stalls the API without needing any malformed input at all."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    for i in range(MAX_PDF_PAGES + 5):
        c.drawString(72, 780, f"Page {i}")
        c.showPage()
    c.save()

    with pytest.raises(ValidationError, match="pages"):
        parse(buf.getvalue(), "pdf")
